"""
Gera pedido_patente_IPS.docx a partir do conteúdo da patente IPS.
Executar: python docs/gerar_patente_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "pedido_patente_IPS.docx")

doc = Document()

# ── Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Estilos base
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(14 if level == 1 else 12)
    if level == 1:
        run.font.color.rgb = RGBColor(0x10, 0x34, 0x6A)
    else:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9A)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def add_bullet(text, highlight=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def add_numbered(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent   = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + " — ")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x10, 0x34, 0x6A)
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    return p

def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '103468')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)

# ════════════════════════════════════════
# CAPA
# ════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("PEDIDO DE PATENTE DE INVENÇÃO")
r.bold = True; r.font.size = Pt(16); r.font.name = 'Arial'
r.font.color.rgb = RGBColor(0x10, 0x34, 0x6A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Instituto Nacional da Propriedade Industrial — INPI")
r2.font.size = Pt(12); r2.font.name = 'Arial'; r2.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Live Equipamentos Ltda.")
r3.bold = True; r3.font.size = Pt(13); r3.font.name = 'Arial'

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Rodrigo Siqueira — Inventor / Depositante")
r4.font.size = Pt(11); r4.font.name = 'Arial'; r4.font.color.rgb = RGBColor(0x44,0x44,0x44)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("Junho / 2026")
r5.font.size = Pt(11); r5.font.name = 'Arial'; r5.font.color.rgb = RGBColor(0x77,0x77,0x77)

doc.add_page_break()

# ════════════════════════════════════════
# 1 — TÍTULO
# ════════════════════════════════════════
add_heading("1 — TÍTULO")
add_hr()
add_body(
    "Sistema Integrado de Análise Biomecânica em Tempo Real para Ambientes de Fitness "
    "com Correção Autônoma por Inteligência Artificial, Gestão Inteligente de Equipamentos "
    "e Monitoramento Preditivo por Sensores IoT"
)

# ════════════════════════════════════════
# 2 — BREVE RELATO
# ════════════════════════════════════════
add_heading("2 — BREVE RELATO")
add_hr()
add_body(
    "O presente invento refere-se a um sistema computacional integrado, denominado IPS "
    "(Intelligent Performance System), que combina visão computacional, inteligência artificial, "
    "processamento de linguagem natural e sensores eletrônicos para realizar, em tempo real, "
    "as seguintes funções em ambientes de fitness, academias, estúdios de Pilates e centros de reabilitação:"
)
add_bullet("(a) Análise biomecânica automatizada do movimento humano durante a prática de exercícios físicos, por meio de câmeras e modelos de estimação de pose corporal com 133 pontos de referência anatômica;")
add_bullet("(b) Detecção automática de erros de execução e emissão de correções sonoras ao praticante em tempo real, sem intervenção humana;")
add_bullet("(c) Registro longitudinal e automático de um prontuário biomecânico individual para cada praticante, desde o primeiro treino, com geração de gráficos e relatórios de evolução objetiva;")
add_bullet("(d) Gestão inteligente do espaço de treino com detecção de ocupação de equipamentos em tempo real e adaptação automática do programa de treino do praticante conforme disponibilidade;")
add_bullet("(e) Monitoramento contínuo da saúde estrutural dos equipamentos de fitness por meio de sensores IoT, com geração de alertas de manutenção preventiva antes da ocorrência de falha.")
add_body("O sistema é aplicável ao setor técnico de equipamentos esportivos com tecnologia embarcada, ciências do movimento humano, gestão de estabelecimentos de fitness e manutenção preditiva de ativos físicos.")

# ════════════════════════════════════════
# 3 — ESTADO DA TÉCNICA
# ════════════════════════════════════════
add_heading("3 — ESTADO DA TÉCNICA")
add_hr()

add_subheading("3.1 Problema a ser resolvido")
add_body("Os estabelecimentos de fitness enfrentam três problemas estruturais não resolvidos pelas soluções disponíveis no mercado:")
add_numbered("Problema 1", "Invisibilidade do progresso: O praticante de academia não dispõe de dados objetivos sobre a qualidade de sua execução nem sobre a evolução real do seu corpo ao longo do tempo. O acompanhamento é feito de forma subjetiva pelo instrutor, limitado pela sua capacidade de atenção simultânea a múltiplos alunos. A ausência de evidência objetiva de progresso é apontada como principal causa de cancelamento de matrículas.")
add_numbered("Problema 2", "Execução incorreta sem correção imediata: Erros posturais durante a prática de exercícios são a principal causa de lesões em academia. O instrutor presente na sala não consegue monitorar todos os alunos em tempo real, especialmente em horários de pico. Nenhum sistema automatizado disponível comercialmente detecta e corrige erros de execução de forma autônoma.")
add_numbered("Problema 3", "Quebra de equipamentos em momentos críticos: A manutenção de equipamentos de fitness é majoritariamente reativa — o aparelho quebra, o aluno reclama, a manutenção é chamada. Não há sistema que monitore continuamente o estado físico dos equipamentos e gere alertas preventivos baseados em dados reais de desgaste.")

add_subheading("3.2 Estado da técnica atual")
add_numbered("Softwares de gestão de academia", "(ex.: Tecnofit, Acade, Club Software, CloudGym): controlam aspectos operacionais e financeiros — cobrança, check-in, agendamento de aulas, controle de acesso. Não possuem qualquer funcionalidade de análise de movimento, monitoramento biomecânico ou avaliação de qualidade de execução.")
add_numbered("Sistemas de fisioterapia remota", "(ex.: Sword Health, Hinge Health, Kaia Health): utilizam câmeras ou sensores para análise de movimento em contexto clínico. São voltados a pacientes com diagnóstico médico específico, em ambiente domiciliar, sem integração com equipamentos de fitness.")
add_numbered("Aplicativos de treino com análise de postura", "(ex.: Kemtai, Onyx): utilizam câmeras de celular para avaliação de postura em exercícios isolados, em ambiente domiciliar. Não integram dados de múltiplos usuários, não se conectam a equipamentos físicos de academia.")

add_subheading("3.3 Inconvenientes do estado da técnica")
add_bullet("Nenhuma solução existente integra análise biomecânica, gestão de espaço, prontuário longitudinal e manutenção preditiva em um único sistema;")
add_bullet("As soluções de análise de movimento existentes são de uso individual e domiciliar, não projetadas para ambientes multiusuário simultâneo;")
add_bullet("Nenhuma solução monitora automaticamente a saúde estrutural dos equipamentos de fitness;")
add_bullet("As soluções clínicas existentes não se aplicam ao contexto de academias e estúdios de fitness;")
add_bullet("Nenhuma solução gera automaticamente um histórico biomecânico longitudinal e individual do praticante sem intervenção manual do instrutor.")

# ════════════════════════════════════════
# 4 — DESCRIÇÃO DO INVENTO
# ════════════════════════════════════════
doc.add_page_break()
add_heading("4 — DESCRIÇÃO DO INVENTO")
add_hr()
add_body("O sistema IPS é composto pelos seguintes blocos funcionais, identificados numericamente conforme o diagrama de blocos apresentado em formato A4 em documento anexo:")

blocos = [
    ("Bloco 1 — Módulo de Captura de Imagem (câmeras)",
     "Constituído por uma ou mais câmeras digitais posicionadas estrategicamente no ambiente de fitness, com campo de visão cobrindo as zonas de exercício. As câmeras capturam sequências de imagens em formato de vídeo digital (frames) a uma taxa mínima de 25 quadros por segundo, em resolução mínima de 1280×720 pixels (HD), transmitindo o fluxo de dados ao Módulo de Processamento (Bloco 2) via interface USB, HDMI ou rede local (Ethernet/Wi-Fi). Para implementações em equipamentos próprios Live Equipamentos (linha V12), as câmeras são integradas fisicamente à estrutura do aparelho."),
    ("Bloco 2 — Módulo de Processamento Central (unidade de computação embarcada)",
     "Constituído por uma unidade de processamento com capacidade de execução de redes neurais profundas em tempo real. Em fase de produção, é utilizado o módulo NVIDIA Jetson Orin NX com 16GB de memória LPDDR5, equipado com acelerador de inteligência artificial com capacidade de 100 TOPS. Em fase de desenvolvimento e piloto, pode ser utilizado computador com GPU dedicada com suporte a CUDA. O Módulo de Processamento Central recebe os frames do Bloco 1, executa o pipeline de análise e coordena todos os demais blocos do sistema."),
    ("Bloco 3 — Módulo de Estimação de Pose (extrator RTMPose)",
     "Software executado sobre o Bloco 2, responsável por identificar e rastrear, em cada frame, as coordenadas espaciais dos 133 pontos de referência anatômica do corpo humano definidos pelo padrão COCO-WholeBody, incluindo articulações de membros superiores, membros inferiores, coluna vertebral, face e extremidades. O extrator utiliza o modelo RTMPose-l, executado via runtime ONNX com aceleração CUDA. Para cada pessoa detectada, o módulo retorna 133 pares de coordenadas normalizadas com índice de confiança por ponto. Suporta detecção e rastreamento simultâneo de múltiplos praticantes."),
    ("Bloco 4 — Módulo de Cálculo de Ângulos Articulares",
     "Software que recebe os vetores de landmarks do Bloco 3 e calcula os ângulos articulares em graus para joelhos, quadris, tornozelos e cotovelos (esquerdo e direito), além de métricas de postura corporal: inclinação do tronco, altura relativa dos ombros, desvio lateral da cabeça, inclinação do pescoço e índice de assimetria bilateral. Os cálculos utilizam produto interno normalizado (arcosseno), com precisão de 0,01°."),
    ("Bloco 5 — Motor de Regras Biomecânicas",
     "Software modular que compara os ângulos calculados pelo Bloco 4 com limites de referência calibrados por exercício e por praticante. Analisa janela temporal de N frames para identificar padrões de erro sustentados, evitando alertas por variações momentâneas. Ao detectar erro, gera objeto de alerta contendo: mensagem de correção em linguagem natural, nível de severidade (alto/médio/baixo) e descrição técnica para log. Implementa mecanismo de debounce que impede repetição do mesmo alerta em intervalo inferior a 4 segundos."),
    ("Bloco 6 — Módulo de Síntese de Voz (Coach Autônomo)",
     "Software que converte os objetos de alerta do Bloco 5 em saída de áudio em linguagem natural, emitida ao praticante em tempo real por alto-falante. A síntese de voz utiliza motor TTS offline de alta qualidade, sem dependência de conexão à internet. Opera em thread dedicada, de forma assíncrona ao loop principal, sem interferir na continuidade do processamento de vídeo."),
    ("Bloco 7 — Módulo de Registro de Sessão e Prontuário Biomecânico",
     "Banco de dados e software de persistência que registra, a cada sessão, o conjunto completo de dados gerados pelos Blocos 3, 4 e 5: landmarks, ângulos articulares, alertas emitidos, tempo de execução e índices de qualidade. Os dados são estruturados por praticante, sessão e exercício, formando um histórico longitudinal denominado prontuário biomecânico individual. Gera relatórios de evolução em gráficos temporais com identificação automática de marcos de progresso."),
    ("Bloco 8 — Módulo de Gestão de Espaço em Tempo Real (Smart Floor)",
     "Software que identifica, a partir dos dados de detecção do Bloco 3, quais zonas pré-configuradas da sala de treino estão ocupadas e por qual praticante. Mantém mapa dinâmico de ocupação atualizado a cada ciclo. Quando o equipamento prescrito ao praticante está ocupado, identifica automaticamente exercícios alternativos com equipamentos disponíveis e notifica o praticante via aplicativo móvel (Bloco 10) em tempo real."),
    ("Bloco 9 — Módulo de Sensores IoT (manutenção preditiva)",
     "Conjunto de sensores eletrônicos instalados nos equipamentos, conectados ao Bloco 2 via Wi-Fi ou Bluetooth Low Energy. Cada unidade sensorial é composta por: microcontrolador ESP32; sensor acelerômetro triaxial (vibração mecânica); sensor de temperatura NTC (monitoramento térmico); sensor de efeito Hall (ciclos de rotação acumulados). Algoritmos de detecção de anomalia comparam leituras com baselines históricos, gerando alertas de manutenção preventiva ao Bloco 11 quando padrões anômalos são detectados."),
    ("Bloco 10 — Aplicativo Móvel do Praticante",
     "Interface de software no dispositivo móvel do praticante, que exibe: programa de treino do dia com equipamentos disponíveis, substituições propostas pelo Bloco 8, feedback de qualidade ao final de cada série e prontuário biomecânico com gráficos de evolução histórica. Comunicação via WebSocket sobre HTTP/HTTPS com o Bloco 2."),
    ("Bloco 11 — Dashboard de Gestão em Tempo Real",
     "Interface web para o gestor do estabelecimento, que consolida em tempo real: mapa de ocupação da sala (Bloco 8), indicadores de presença e circulação de funcionários, alertas de manutenção preventiva (Bloco 9), relatórios de evolução biomecânica individual e agregada. Para redes com múltiplas unidades, apresenta visão centralizada de todas as filiais a partir de uma única interface."),
]

for titulo, descricao in blocos:
    add_subheading(titulo)
    add_body(descricao)

# ════════════════════════════════════════
# 5 — FUNCIONAMENTO
# ════════════════════════════════════════
doc.add_page_break()
add_heading("5 — FUNCIONAMENTO DO INVENTO")
add_hr()
add_body("O funcionamento do sistema IPS ocorre de forma cíclica e contínua durante o período de operação do estabelecimento, conforme descrito a seguir:")

etapas = [
    ("Etapa 1 — Inicialização e calibração",
     "O Módulo de Processamento Central (Bloco 2) carrega na memória o modelo RTMPose-l e o mapeia para a unidade de aceleração de IA disponível (GPU ou NPU). As câmeras do Bloco 1 são inicializadas e o fluxo de vídeo é estabelecido. Os sensores IoT (Bloco 9) estabelecem conexão Wi-Fi e iniciam transmissão de dados. O sistema carrega os perfis de praticantes ativos e os programas de treino associados."),
    ("Etapa 2 — Captura e extração de pose",
     "A cada frame capturado pelas câmeras (Bloco 1), o Módulo de Estimação de Pose (Bloco 3) executa o modelo RTMPose-l, identificando todas as pessoas presentes no campo de visão. Para cada pessoa detectada, são retornadas as coordenadas dos 133 pontos anatômicos com seus respectivos índices de confiança. Pontos com confiança inferior a 0,3 (30%) são descartados dos cálculos subsequentes."),
    ("Etapa 3 — Cálculo de ângulos e avaliação biomecânica",
     "Os landmarks extraídos no Bloco 3 são encaminhados ao Módulo de Cálculo de Ângulos (Bloco 4), que computa os ângulos articulares e as métricas de postura para cada praticante. Os resultados são encaminhados ao Motor de Regras Biomecânicas (Bloco 5), que avalia a execução do exercício vigente. Se um erro de execução é confirmado por N frames consecutivos (tipicamente N=15 a 30 a 30fps), o motor emite um objeto de alerta."),
    ("Etapa 4 — Correção autônoma por voz",
     "O objeto de alerta gerado pelo Bloco 5 é recebido pelo Módulo de Síntese de Voz (Bloco 6), que converte a mensagem de correção em áudio e o emite ao praticante em tempo real. O sistema verifica o mecanismo de debounce antes da emissão para evitar repetição excessiva da mesma correção."),
    ("Etapa 5 — Registro no prontuário",
     "Em paralelo aos ciclos de análise, o Módulo de Registro (Bloco 7) persiste continuamente os dados de ângulos, alertas e métricas de execução na base de dados, associados ao praticante identificado e à sessão em curso. Ao término da sessão, o módulo consolida os dados e atualiza o prontuário biomecânico individual."),
    ("Etapa 6 — Gestão de espaço e adaptação de treino",
     "Simultaneamente ao ciclo de análise biomecânica, o Módulo Smart Floor (Bloco 8) atualiza o mapa de ocupação da sala. Quando detecta que um equipamento prescrito ao praticante está ocupado, o módulo consulta o programa de treino, seleciona exercício alternativo com equipamento disponível e notifica o praticante via aplicativo móvel (Bloco 10) em tempo real."),
    ("Etapa 7 — Monitoramento preditivo de equipamentos",
     "De forma contínua e paralela, os sensores do Bloco 9 transmitem leituras de vibração, temperatura e ciclos de uso ao processador central. O algoritmo de detecção de anomalia compara as leituras com o baseline histórico do equipamento. Ao identificar padrão anômalo indicativo de desgaste, o sistema registra um alerta de manutenção preventiva no Dashboard (Bloco 11), informando o gestor com antecedência suficiente para agendar a intervenção sem impacto na operação."),
    ("Etapa 8 — Visualização e gestão",
     "O Dashboard (Bloco 11) consolida em tempo real todos os dados gerados pelos blocos anteriores, apresentando ao gestor uma visão completa do estado do estabelecimento: praticantes em atividade, ocupação por equipamento, indicadores de funcionários, alertas de manutenção e relatórios de evolução biomecânica individual e agregada."),
]

for titulo, texto in etapas:
    add_numbered(titulo, texto)

# ════════════════════════════════════════
# 6 — VANTAGENS
# ════════════════════════════════════════
doc.add_page_break()
add_heading("6 — VANTAGENS")
add_hr()

vantagens = [
    "Primeiro sistema no Brasil a integrar análise biomecânica, gestão de espaço, prontuário longitudinal e manutenção preditiva em um único produto, eliminando a necessidade de múltiplos sistemas desconexos.",
    "Correção autônoma em tempo real sem dependência de instrutor presente: o sistema monitora simultaneamente todos os praticantes na sala, sem limitação de atenção humana, garantindo a mesma qualidade de supervisão em qualquer horário.",
    "Prontuário biomecânico objetivo e automático: o praticante dispõe de evidência científica e objetiva de sua evolução física, gerada automaticamente sem intervenção manual, reduzindo significativamente a taxa de cancelamento de matrículas.",
    "Compatibilidade com equipamentos já existentes na academia: o sistema utiliza visão computacional para análise de movimento, não requerendo substituição dos aparelhos presentes no estabelecimento, reduzindo custo e tempo de implantação.",
    "Gestão inteligente de espaço com zero tempo de espera: a adaptação automática do treino conforme disponibilidade de equipamentos em tempo real elimina o tempo improdutivo do praticante aguardando aparelhos.",
    "Manutenção preditiva com redução de custos operacionais: a detecção antecipada de desgaste em equipamentos permite agendamento planejado de manutenção, evitando quebras em horários de pico e reduzindo custos de manutenção emergencial.",
    "Escalabilidade para redes e franquias: o sistema opera com visão centralizada de múltiplas unidades a partir de uma única plataforma, permitindo padronização de atendimento e qualidade de operação em todas as filiais simultaneamente.",
    "Funcionamento offline e de baixa latência: o processamento é executado localmente na unidade embarcada (Bloco 2), sem dependência de conexão à internet para as funções de análise biomecânica e correção de voz.",
    "Dados longitudinais para prevenção de lesões: o histórico biomecânico individual permite identificar padrões de assimetria e sobrecarga articular ao longo do tempo, possibilitando intervenção preventiva antes do desenvolvimento de lesão.",
    "Modelo de receita recorrente escalável: a combinação de hardware embarcado com assinatura de software cria receita previsível e recorrente, aplicável tanto a academias individuais quanto a redes com centenas de unidades.",
]

for i, v in enumerate(vantagens, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent   = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{i}. ")
    r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x10, 0x34, 0x6A)
    r2 = p.add_run(v)
    r2.font.name = 'Arial'; r2.font.size = Pt(11)

# ── Rodapé
doc.add_paragraph()
add_hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Live Equipamentos Ltda. | Rodrigo Siqueira — Inventor / Depositante | Junho 2026")
r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

doc.save(OUT)
print(f"Salvo: {OUT}")

"""
Gera Descritivo_TCO_IPS.docx — narrativa explicativa do TCO para o investidor.
Executar: python docs/gerar_descritivo_tco.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "Descritivo_TCO_IPS.docx")

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.2)
    section.right_margin  = Cm(2.5)

AZUL   = RGBColor(0x10, 0x34, 0x6A)
VERDE  = RGBColor(0x1B, 0x5E, 0x20)
GOLD   = RGBColor(0xC6, 0x86, 0x00)
CINZA  = RGBColor(0x44, 0x44, 0x44)
PRETO  = RGBColor(0x11, 0x11, 0x11)

def hr(cor="103468"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), cor)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)

def titulo(text, size=15, color=AZUL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def secao(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{num}  ")
    r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(13)
    r1.font.color.rgb = GOLD
    r2 = p.add_run(text.upper())
    r2.bold = True; r2.font.name = 'Arial'; r2.font.size = Pt(13)
    r2.font.color.rgb = AZUL
    return p

def corpo(text, justify=True, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = PRETO
    return p

def corpo_bold(parts, justify=True, space_after=8):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.color.rgb = PRETO if not bold else AZUL
    return p

def destaque(text, color=VERDE, size=12, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(11)
        r1.font.color.rgb = AZUL
    r2 = p.add_run(("— " if bold_prefix else "• ") + text)
    r2.font.name = 'Arial'; r2.font.size = Pt(11)
    r2.font.color.rgb = PRETO
    return p

def caixa_numero(label, valor, sublabel, color=AZUL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{label}:  ")
    r1.font.name = 'Arial'; r1.font.size = Pt(11); r1.font.color.rgb = CINZA
    r2 = p.add_run(valor)
    r2.bold = True; r2.font.name = 'Arial'; r2.font.size = Pt(13)
    r2.font.color.rgb = color
    r3 = p.add_run(f"  {sublabel}")
    r3.font.name = 'Arial'; r3.font.size = Pt(10); r3.font.color.rgb = CINZA

# ═══════════════════════════════════════
# CAPA
# ═══════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("IPS PLATFORM")
r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(28)
r.font.color.rgb = AZUL

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Análise de Custo Total de Propriedade")
r2.bold = True; r2.font.name = 'Arial'; r2.font.size = Pt(16)
r2.font.color.rgb = CINZA

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Modelo Edge Pesado  vs.  Modelo Data Center Híbrido")
r3.font.name = 'Arial'; r3.font.size = Pt(13); r3.font.color.rgb = GOLD; r3.bold = True

doc.add_paragraph()

for linha in ["100 academias  |  Horizonte de 36 meses",
              "Live Equipamentos Ltda.  |  Rodrigo Siqueira",
              "Junho / 2026  |  Confidencial"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(linha)
    r.font.name = 'Arial'; r.font.size = Pt(11); r.font.color.rgb = CINZA

doc.add_page_break()

# ═══════════════════════════════════════
# 1. CONTEXTO
# ═══════════════════════════════════════
secao("01", "Contexto e objetivo da análise")
hr()

corpo(
    "O IPS (Intelligent Performance System) é uma plataforma de análise biomecânica em tempo real "
    "desenvolvida pela Live Equipamentos para ambientes de fitness. O sistema utiliza câmeras e "
    "inteligência artificial para rastrear 133 pontos do corpo humano, detectar erros de execução "
    "de exercícios e emitir correções por voz ao praticante — sem intervenção humana."
)
corpo(
    "Desde sua concepção, o sistema foi projetado para operar de forma totalmente local (edge computing), "
    "com todo o processamento de IA acontecendo dentro do próprio equipamento, em hardware embarcado "
    "de alto desempenho (NVIDIA Jetson Orin NX). Essa arquitetura garante baixíssima latência e "
    "funcionamento sem dependência de internet."
)
corpo(
    "No entanto, à medida que o produto se prepara para escalar de unidades piloto para centenas de "
    "academias, surge uma questão estratégica fundamental: "
    "vale a pena manter todo o processamento no aparelho (Modelo A) "
    "ou centralizar o poder computacional em um data center próprio e reduzir drasticamente "
    "o custo do hardware instalado em cada academia (Modelo B)?"
)
corpo(
    "Este documento apresenta a análise comparativa de Custo Total de Propriedade (TCO) "
    "entre as duas arquiteturas para um cenário de 100 academias ao longo de 36 meses, "
    "e demonstra o impacto financeiro direto de cada escolha na margem do negócio."
)

# ═══════════════════════════════════════
# 2. OS DOIS MODELOS
# ═══════════════════════════════════════
secao("02", "Os dois modelos arquiteturais")
hr()

corpo_bold([
    ("Modelo A — Edge Pesado (arquitetura atual):", True),
    (" cada academia recebe um equipamento completo com processador NVIDIA Jetson Orin NX "
     "(R$ 5.800 por unidade), capaz de executar o modelo de inteligência artificial inteiramente "
     "de forma local. O sistema não depende de internet para funcionar. O custo de hardware "
     "por academia é de ", False),
    ("R$ 8.600", True),
    (", incluindo câmeras, cabeamento e instalação.", False),
], space_after=12)

corpo_bold([
    ("Modelo B — Data Center Híbrido (proposta de evolução):", True),
    (" cada academia recebe apenas um dispositivo de borda leve e barato "
     "(Raspberry Pi 5 com chip de IA dedicado HAILO-8, ~R$ 1.200), "
     "responsável somente por capturar as imagens e extrair os pontos do corpo "
     "(133 coordenadas por frame). Esses dados — extremamente leves, sem imagem — "
     "são enviados ao data center central, que realiza toda a análise biomecânica, "
     "aplica as regras de correção e devolve o resultado em menos de 150 milissegundos. "
     "O custo de hardware por academia cai para ", False),
    ("R$ 3.800", True),
    (" — uma redução de 56%.", False),
], space_after=12)

destaque(
    "A diferenca critica: no Modelo B, um unico data center serve 100 academias simultaneamente. "
    "Em vez de 100 processadores de R$ 5.800, compra-se 2 servidores de R$ 10.000 cada.",
    color=AZUL, size=11
)

# ═══════════════════════════════════════
# 3. HARDWARE
# ═══════════════════════════════════════
secao("03", "Comparativo de hardware por academia")
hr()

corpo(
    "A tabela a seguir resume os componentes de hardware instalados em cada academia "
    "nos dois modelos, com seus respectivos custos:"
)

# Tabela simples como texto estruturado
for linha in [
    ("Componente",               "Modelo A (Edge)",           "Modelo B (Data Center)"),
    ("Processador de borda",     "Jetson Orin NX — R$ 5.800", "Raspberry Pi 5 + HAILO-8 — R$ 1.200"),
    ("Cameras HD (3 unidades)",  "R$ 1.200",                  "R$ 1.200"),
    ("Cabeamento e instalacao",  "R$ 800",                    "R$ 600"),
    ("Display/tablet do aluno",  "R$ 800",                    "R$ 800"),
    ("TOTAL POR ACADEMIA",       "R$ 8.600",                  "R$ 3.800"),
    ("TOTAL 100 ACADEMIAS",      "R$ 860.000",                "R$ 380.000"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    is_header = linha[0] in ("Componente",)
    is_total  = "TOTAL" in linha[0]
    tab = p.add_run(f"{linha[0]:<32}  {linha[1]:<32}  {linha[2]}")
    tab.font.name = 'Courier New'
    tab.font.size = Pt(9.5)
    tab.bold = is_total or is_header
    tab.font.color.rgb = AZUL if is_total else (CINZA if is_header else PRETO)

doc.add_paragraph()

corpo_bold([
    ("Economia em hardware por academia: ", False),
    ("R$ 4.800 (56% menor).", True),
    (" Para 100 academias, a diferença de hardware é de ", False),
    ("R$ 480.000 ", True),
    ("— dinheiro que deixa de ser imobilizado em processadores espalhados pelas academias "
     "e passa a ser centralizado em infraestrutura gerenciável e atualizável remotamente.", False),
])

# ═══════════════════════════════════════
# 4. DATA CENTER
# ═══════════════════════════════════════
secao("04", "Investimento no data center")
hr()

corpo(
    "O Modelo B requer um investimento inicial em infraestrutura central. "
    "Para atender 100 academias simultaneamente, o setup necessário é:"
)
bullet("2 servidores dedicados (processamento de regras e armazenamento) — R$ 20.000")
bullet("UPS 3kVA, rack, switches e cabeamento interno — R$ 10.000")
bullet("Instalacao, configuracao e hardening de segurança — R$ 5.000")

destaque("Investimento total no data center: R$ 35.000", color=VERDE, size=12)

corpo(
    "Este valor — R$ 35.000 — representa apenas 4% da economia gerada em hardware "
    "(R$ 480.000). Ou seja: o data center se paga com a reducao de hardware das primeiras "
    "7 academias instaladas."
)

corpo(
    "Mensalmente, os custos operacionais do data center (colocation, link dedicado de 100Mbps, "
    "storage e backup) somam aproximadamente R$ 5.800. Dividido por 100 academias, "
    "representa R$ 58 por academia por mês — custo marginal irrisório."
)

corpo(
    "Uma característica importante deste modelo: o processamento de keypoints (coordenadas do corpo) "
    "não requer GPU de alto desempenho no data center. O dispositivo de borda na academia realiza "
    "a parte pesada da visão computacional localmente. O data center recebe apenas dados "
    "estruturados — cerca de 2KB por frame — e executa regras e cálculos de baixo custo computacional. "
    "Isso mantém a infraestrutura central simples, barata e escalável."
)

# ═══════════════════════════════════════
# 5. TCO TOTAL
# ═══════════════════════════════════════
secao("05", "TCO total — 36 meses, 100 academias")
hr()

corpo(
    "Consolidando todos os custos — hardware instalado nas academias, infraestrutura do "
    "data center e custos mensais operacionais ao longo de 36 meses — chegamos ao "
    "seguinte comparativo:"
)

doc.add_paragraph()
caixa_numero("TCO Modelo A (Edge Pesado)",    "R$ 1.436.000", "em 36 meses",  color=RGBColor(0xC6, 0x28, 0x28))
caixa_numero("TCO Modelo B (Data Center)",    "R$   911.800", "em 36 meses",  color=VERDE)
caixa_numero("Economia total",                "R$   524.200", "(36% menor)",   color=AZUL)
caixa_numero("TCO por academia / mes (A)",    "R$       399", "por academia",  color=RGBColor(0xC6, 0x28, 0x28))
caixa_numero("TCO por academia / mes (B)",    "R$       253", "por academia",  color=VERDE)
doc.add_paragraph()

corpo(
    "O Modelo B é R$ 524.200 mais barato ao longo de 3 anos — mesmo considerando "
    "o investimento inicial no data center. Essa economia se explica principalmente "
    "pela redução do custo de hardware por academia (de R$ 8.600 para R$ 3.800) "
    "e pelo menor custo de suporte técnico, já que atualizações de software e do "
    "modelo de IA passam a ser feitas centralmente, sem visitas técnicas a cada unidade."
)

# ═══════════════════════════════════════
# 6. IMPACTO NA MARGEM
# ═══════════════════════════════════════
secao("06", "Impacto na margem do negocio")
hr()

corpo(
    "O efeito mais relevante do Modelo B não está apenas na redução de custo — "
    "está na transformação do modelo de negócio. "
    "Considerando uma mensalidade SaaS de R$ 400 por academia:"
)

doc.add_paragraph()
for linha, val_a, val_b, color_b in [
    ("Receita mensal por academia",  "R$ 400",  "R$ 400",  PRETO),
    ("Custo por academia / mes",     "R$ 399",  "R$ 253",  PRETO),
    ("MARGEM BRUTA por academia",    "R$ 1",    "R$ 147",  VERDE),
    ("MARGEM 100 academias / mes",   "R$ 100",  "R$ 14.700", VERDE),
    ("MARGEM BRUTA ANUAL",           "R$ 1.200","R$ 176.400", VERDE),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{linha:<38}")
    r1.font.name = 'Courier New'; r1.font.size = Pt(10)
    r1.font.color.rgb = CINZA if "MARGEM" not in linha else AZUL
    r1.bold = "MARGEM" in linha
    r2 = p.add_run(f"  {val_a:<16}  {val_b}")
    r2.font.name = 'Courier New'; r2.font.size = Pt(10)
    r2.bold = "MARGEM" in linha
    r2.font.color.rgb = color_b if "MARGEM" in linha else PRETO

doc.add_paragraph()

corpo(
    "No Modelo A, cada academia gera R$ 1 de margem por mês. "
    "O negócio opera essencialmente no break-even — toda a receita é consumida "
    "pelo custo de hardware amortizado e suporte."
)
corpo(
    "No Modelo B, a margem por academia sobe para R$ 147 por mês. "
    "Com 100 academias, isso gera R$ 14.700 de margem bruta mensal — "
    "ou R$ 176.400 por ano. O data center converte um negócio de hardware em "
    "um negócio SaaS com margem real e recorrente."
)

destaque(
    "A margem sobe de R$ 1 para R$ 147 por academia/mes — um aumento de 14.600%.",
    color=VERDE, size=12
)

# ═══════════════════════════════════════
# 7. VANTAGENS ADICIONAIS
# ═══════════════════════════════════════
secao("07", "Vantagens estrategicas alem do custo")
hr()

corpo("Alem do impacto financeiro direto, o Modelo B traz vantagens estrategicas significativas:")

bullet(
    "Atualizacao centralizada do modelo de IA: melhorias no algoritmo de analise biomecanica "
    "sao implantadas em todas as academias simultaneamente, sem necessidade de visita tecnica "
    "ou troca de hardware. No Modelo A, uma atualizacao em 100 unidades exige logistica complexa.",
    bold_prefix="Agilidade tecnologica"
)
bullet(
    "Todos os dados biomecanicos de praticantes ficam centralizados em infraestrutura propria, "
    "facilitando conformidade com a LGPD, geracao de relatorios agregados e "
    "construcao do banco de dados biomecanico que representa o principal ativo "
    "de longo prazo da empresa para uma eventual saida estrategica (M&A ou IPO).",
    bold_prefix="Centralizacao de dados"
)
bullet(
    "Com hardware mais barato na ponta, o ticket de entrada para a academia cai, "
    "acelerando a adocao. A empresa pode oferecer o hardware em comodato ou leasing "
    "a custo muito menor, removendo a barreira de capex para o cliente.",
    bold_prefix="Menor barreira de adocao"
)
bullet(
    "O mesmo data center que serve 100 academias pode servir 500 com apenas "
    "um upgrade de servidores. O crescimento nao exige proporcionalidade "
    "em investimento de infraestrutura.",
    bold_prefix="Escalabilidade"
)
bullet(
    "Monitoramento centralizado de todas as unidades em tempo real — "
    "identifica falhas de hardware, quedas de desempenho e anomalias em qualquer academia "
    "imediatamente, sem depender de relato do cliente.",
    bold_prefix="Visibilidade operacional"
)

# ═══════════════════════════════════════
# 8. PONTOS DE ATENÇÃO
# ═══════════════════════════════════════
secao("08", "Pontos de atencao e mitigacoes")
hr()

corpo("O Modelo B apresenta dois pontos que requerem atencao e mitigacao adequada:")

corpo_bold([("Dependencia de internet:", True),
            (" ao contrario do Modelo A, o Modelo B requer conectividade para as funcoes "
             "de analise em tempo real. Mitigation: o dispositivo de borda mantem "
             "um modo offline basico (contagem de presenca e alerta de postura simplificado) "
             "para os momentos de instabilidade de rede. A grande maioria das academias "
             "em areas urbanas ja opera com link dedicado.", False)])

corpo_bold([("Latencia de rede:", True),
            (" a correcao por voz precisa chegar ao praticante em menos de 400-500ms "
             "para ser percebida como sincrona com o movimento. Testes indicam latencia "
             "de 50-150ms com link de 10Mbps — dentro do limite. O ponto critico e "
             "validar em campo com a qualidade de internet tipica das academias-alvo.", False)])

# ═══════════════════════════════════════
# 9. CONCLUSÃO
# ═══════════════════════════════════════
secao("09", "Conclusao e recomendacao")
hr()

corpo(
    "A analise demonstra que o investimento em infraestrutura de data center propria "
    "e financeiramente justificado e estrategicamente necessario para a escala do IPS Platform. "
    "O custo de montagem da infraestrutura central (R$ 35.000) representa apenas 4% "
    "da economia gerada em hardware nas primeiras 100 academias (R$ 480.000)."
)
corpo(
    "Mais do que uma questao de custo, o data center e o que permite transformar "
    "o IPS de um produto de hardware com margem minima em uma plataforma SaaS "
    "com margem bruta real, crescimento recorrente e ativo de dados valioso "
    "para uma saida estrategica."
)

destaque(
    "Recomendacao: investir na infraestrutura de data center como parte do "
    "round de investimento, com alocacao estimada de R$ 35.000 a R$ 50.000 "
    "para o setup inicial, cobrindo 100 academias com capacidade de expansao "
    "para 300-400 unidades com o mesmo hardware.",
    color=AZUL, size=11
)

# Rodape
doc.add_paragraph()
hr("CCCCCC")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Live Equipamentos Ltda.  |  Rodrigo Siqueira  |  Junho 2026  |  Documento Confidencial"
)
r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print(f"Salvo: {OUT}")
